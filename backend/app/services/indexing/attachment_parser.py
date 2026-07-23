"""
Attachment text extractor — supports PDF (pypdf), DOCX (python-docx), and plain text.
Falls back gracefully to stub messages if optional dependencies are missing.
"""
import os
import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class AttachmentParser:
    """Parses text content from email attachments (PDF, DOCX, TXT, CSV, Markdown)."""

    # ─── PDF ──────────────────────────────────────────────────────────────────
    def _extract_pdf(self, content_bytes: bytes, filename: str) -> str:
        try:
            import pypdf  # type: ignore
            reader = pypdf.PdfReader(io.BytesIO(content_bytes))
            pages_text = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text.strip())
            extracted = "\n\n".join(pages_text)
            if extracted.strip():
                logger.info(f"[PDF] Extracted {len(extracted)} chars from {filename}")
                return extracted
            return f"[PDF: {filename} — no extractable text, may be a scanned image]"
        except ImportError:
            logger.warning("pypdf not installed — using stub for PDF extraction")
            return f"[PDF Document: {filename}]"
        except Exception as e:
            logger.warning(f"PDF extraction failed for {filename}: {e}")
            return f"[PDF: {filename} — extraction error]"

    # ─── DOCX ─────────────────────────────────────────────────────────────────
    def _extract_docx(self, content_bytes: bytes, filename: str) -> str:
        try:
            from docx import Document  # type: ignore
            doc = Document(io.BytesIO(content_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also capture table cell content
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            extracted = "\n".join(paragraphs)
            if extracted.strip():
                logger.info(f"[DOCX] Extracted {len(extracted)} chars from {filename}")
                return extracted
            return f"[Word Document: {filename} — no text content found]"
        except ImportError:
            logger.warning("python-docx not installed — using stub for DOCX extraction")
            return f"[Word Document: {filename}]"
        except Exception as e:
            logger.warning(f"DOCX extraction failed for {filename}: {e}")
            return f"[DOCX: {filename} — extraction error]"

    # ─── Plain text / CSV / Markdown / JSON / HTML ────────────────────────────
    def _extract_text(self, content_bytes: bytes, filename: str) -> str:
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                text = content_bytes.decode(encoding)
                logger.info(f"[TEXT] Extracted {len(text)} chars from {filename} ({encoding})")
                return text
            except (UnicodeDecodeError, LookupError):
                continue
        return f"[Text file: {filename} — could not decode]"

    # ─── Images (placeholder — OCR would go here) ─────────────────────────────
    def _extract_image(self, content_bytes: bytes, filename: str) -> str:
        return f"[Image Attachment: {filename} — OCR not enabled in this build]"

    # ─── Public API ────────────────────────────────────────────────────────────
    def extract_text(self, filename: str, content_bytes: bytes, mime_type: str = "") -> str:
        """
        Dispatch to the appropriate extractor based on file extension.
        Always returns a non-empty string (may be a descriptive stub on failure).
        """
        if not content_bytes:
            return f"[Empty attachment: {filename}]"

        ext = os.path.splitext(filename)[1].lower()

        if ext == ".pdf" or mime_type == "application/pdf":
            return self._extract_pdf(content_bytes, filename)
        elif ext in (".docx",) or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self._extract_docx(content_bytes, filename)
        elif ext == ".doc" or mime_type == "application/msword":
            # Legacy .doc — not supported without additional libs
            return f"[Legacy Word .doc: {filename} — convert to .docx for full text extraction]"
        elif ext in (".txt", ".md", ".csv", ".json", ".html", ".htm", ".xml", ".yaml", ".yml", ".log"):
            return self._extract_text(content_bytes, filename)
        elif ext in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".tiff"):
            return self._extract_image(content_bytes, filename)
        else:
            # Generic fallback — try plain text decode
            try:
                decoded = content_bytes.decode("utf-8", errors="ignore")
                if len(decoded.strip()) > 20:
                    return decoded
            except Exception:
                pass
            return f"[Binary Attachment: {filename} ({ext or 'unknown type'})]"


attachment_parser = AttachmentParser()
